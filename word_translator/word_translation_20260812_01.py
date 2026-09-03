r"""
VERSION 2026.0812.01
Word Gemini 翻訳ツール（書式完全保持版）

■ このバージョンでの変更（Gemini APIプロキシ対応）
会社PCからGemini APIへの直接アクセスが遮断された事象（2026-08-10頃）への対応。
旧SDK google.generativeai を直接呼ぶ方式をやめ、共通モジュール
gemini_client.py の generate_advanced() 経由（直接呼び出しが失敗したら
自宅PCのプロキシへ自動フォールバック）へ移行した。
rtocs_organizer / analog_ic_se_strategy_organizer / outlook_total_organizer /
excel_translation / pdf_translator / ppt_translation と同じ方式。

変更した関数は次の3つだけで、Word処理側（段落・表の走査、run単位の書き戻し、
進捗表示、ファイル選択）は一切変更していない。
  `check_dependencies` … google-generativeai のチェックを共通モジュールのチェックへ
  `init_gemini`        … genai.list_models() による自動モデル検出を廃止（下記）
  `translate_batch_gemini` … 互換シム経由の呼び出しへ置換

【重要】自動モデル検出を廃止した理由:
旧版は init_gemini() の中で genai.list_models() を呼んで使用可能モデルを
自動検出していたが、これはネットワークアクセスを伴うため、直接アクセスが
遮断された環境では必ず例外になり「API初期化エラー」で sys.exit(1) していた
（＝プロキシ経由なら翻訳できるのに、ツールが起動すらできない）。
共通モジュール・プロキシのどちらにも list_models 相当が無いため、固定モデル名
（環境変数 GEMINI_MODEL で上書き可。既定 gemini-2.5-flash）を使う方式に変更した。
副作用として、指定モデルが使えない環境では404が出るようになる。

【挙動差】旧版はタイムアウトを指定していなかったため旧SDKの既定に従っていたが、
移行後は gemini_client.py 側の固定値（直接15秒 / プロキシ60秒）になる。呼び出し側
から指定する口は無い。遮断下では最初のバッチだけ直接呼び出しの15秒を待つぶん
遅くなるが、一度失敗すると以降はプロキシ直行になるため2バッチ目以降は影響しない。
これは仕様どおりの挙動であり、不具合ではない。
なお本ツールにはリトライ機構が無いため（旧版からの仕様）、失敗したバッチは
その場で原文のまま返る。これは移行で新たに生じたものではない。

■ 必要な環境変数（どちらか一方以上）
  GEMINI_API_KEY    … 直接呼び出し用
  GEMINI_PROXY_URL  … 自宅PCプロキシのURL（直接呼び出し失敗時のフォールバック先）
  GEMINI_MODEL      … 使用モデルを変えたい場合のみ（任意。既定 gemini-2.5-flash）
  GEMINI_COMMON_DIR … gemini_client.py の置き場所を明示したい場合のみ（任意）

■ 使い方
  1. run_word_translator.bat を実行（または python word_translation_20260812_01.py）
  2. 翻訳したい .docx を選ぶ
  3. 翻訳先言語を選んで「翻訳開始」を押す
  4. 完了すると同じフォルダに `元ファイル名_gemini_japanese.docx` のように保存される

■ 既知の制限（20260306_01 から変わっていない既存仕様）
  - 翻訳対象は本文の段落と表のセルのみ。ヘッダー/フッター・脚注・テキストボックス内の
    文字は翻訳されない。
  - 失敗時のリトライ機構が無い。通信に失敗したバッチはその場で原文のまま残る。
  - ログファイルは作らない。エラーはコンソールへ表示されるだけ。
  - 翻訳中の進捗バーはほとんど動かない（0 → 完了 の2段階）。
"""
import tkinter as tk
from tkinter import filedialog, messagebox
import os
import sys
import shutil
import re
import time
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed

# ============================================================
# Gemini 共通クライアント(gemini_client.py)への互換シム
# ============================================================
# 会社PCからGemini APIへの直接アクセスが遮断される事象(2026-08-10頃)を受け、
# rtocs_organizer / analog_ic_se_strategy_organizer / outlook_total_organizer /
# excel_translation / pdf_translator / ppt_translation と同様に、共通モジュール
# gemini_client.py の
# generate_advanced() 経由(直接呼び出しが失敗したら自宅PCプロキシへ自動フォール
# バック)へ移行した。
#
# 本ツールは旧SDK(google.generativeai)の genai.GenerativeModel(...).generate_content()
# を1箇所で使っていただけだが、他ツールと実装を揃えるため、同じ形の薄い互換シム
# (_CommonGeminiClient)を用意し、そこを経由する方式にした。これにより、レスポンスを
# 読む側(response.parts で空判定 → response.text を正規表現で番号付きリストへ戻す処理)や
# 並列処理・進捗表示のロジックは一切変更しないで済んでいる。
# 旧SDK(google-generativeai)への依存は本バージョンで無くなった。
#
# 必要な環境変数(会社PC):
#   GEMINI_API_KEY   … 直接呼び出し用(gemini_client.py 側が読む)
#   GEMINI_PROXY_URL … 自宅PCプロキシのURL(直接呼び出し失敗時のフォールバック先)
#   GEMINI_MODEL     … 使用モデルを変えたい場合のみ(任意。既定 gemini-2.5-flash)
#   GEMINI_COMMON_DIR… gemini_client.py の置き場所を明示したい場合のみ(任意)

_GEMINI_COMMON_DIR_ENV = os.environ.get("GEMINI_COMMON_DIR")
if _GEMINI_COMMON_DIR_ENV:
    _COMMON_DIR_CANDIDATES = [_GEMINI_COMMON_DIR_ENV]
else:
    # 会社PCでは本スクリプトが PythonScripts\word\word_translator\ に、
    # gemini_client.py が PythonScripts\common\ に置かれるため、正解は「1つ上」では
    # なく「2つ上」の common になる。他ツール(1つ上が common)と同じ配置に置かれた
    # 場合でも動くよう、上位ディレクトリを順に探して最初に見つかったものを使う。
    _SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
    _COMMON_DIR_CANDIDATES = [
        os.path.abspath(os.path.join(_SCRIPT_DIR, *([os.pardir] * _n + ["common"])))
        for _n in (1, 2, 3)
    ]

_COMMON_DIR = next(
    (_d for _d in _COMMON_DIR_CANDIDATES
     if os.path.isfile(os.path.join(_d, "gemini_client.py"))),
    _COMMON_DIR_CANDIDATES[0])

if _COMMON_DIR not in sys.path:
    sys.path.insert(0, _COMMON_DIR)

# gemini_client のインポートはここで試みるが、失敗しても import 時点では落とさない
# (原因が分かるメッセージを起動時チェックで出すため)。
try:
    from gemini_client import generate_advanced as _generate_advanced
    _GEMINI_CLIENT_IMPORT_ERROR = None
except Exception as _e:
    _generate_advanced = None
    _GEMINI_CLIENT_IMPORT_ERROR = _e

# 本ツールは全機能が翻訳(AI呼び出し)のため、共通モジュールを読み込めたかどうかが
# そのまま「Gemini が使えるか」になる。
HAS_GEMINI = _generate_advanced is not None
if not HAS_GEMINI:
    print(f"警告: Gemini共通モジュール(gemini_client.py)を読み込めませんでした: "
          f"{_GEMINI_CLIENT_IMPORT_ERROR}")

# 旧版は genai.list_models() で使用可能モデルを自動検出していたが、これはネットワーク
# アクセスを伴うため、直接アクセスが遮断された環境では必ず失敗し、起動時の
# init_gemini() が False を返して sys.exit(1) していた(＝ツールが起動できない)。
# 共通モジュール・プロキシのどちらにも list_models 相当が無いため、自動検出は廃止し、
# 固定モデル名(環境変数 GEMINI_MODEL で上書き可)を使う方式に変更した。
GEMINI_MODEL_NAME = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")


def _gemini_common_module_error_message():
    """共通モジュールを読み込めなかったときの、原因が分かる案内文を組み立てる。"""
    return ("Gemini共通モジュール(gemini_client.py)を読み込めませんでした。\n"
            f"探索したパス: {' / '.join(_COMMON_DIR_CANDIDATES)}\n"
            f"元のエラー: {_GEMINI_CLIENT_IMPORT_ERROR}\n\n"
            "gemini-common-tools を配置し、必要なら環境変数 GEMINI_COMMON_DIR で\n"
            "gemini_client.py のあるフォルダを指定してください。")


def _schema_to_jsonable(value):
    """REST APIのpayloadへそのまま載せられる素のdict/listへ変換する。
    本ツールの safety_settings は既に素のdictのリストなのでそのまま返るが、
    他ツールのシムと実装を揃えるために残してある(pydanticモデル等が渡された
    場合にJSON化できなくなるのを防ぐ保険)。"""
    if value is None or isinstance(value, (dict, list, str, int, float, bool)):
        return value
    # pydantic v2 (model_dump) / v1 (dict) の両方に対応。REST APIのフィールド名は
    # camelCase なので by_alias=True で別名を使う。
    for attr, kwargs in (("model_dump", {"mode": "json", "exclude_none": True, "by_alias": True}),
                         ("dict", {"exclude_none": True, "by_alias": True})):
        fn = getattr(value, attr, None)
        if callable(fn):
            try:
                return fn(**kwargs)
            except Exception:
                try:
                    return fn()
                except Exception:
                    pass
    return value


class _GeminiGenerateConfig:
    """旧SDKの genai.types.GenerationConfig(...) ＋ safety_settings 相当の設定
    オブジェクト。旧SDKへの依存を断つため、同等の入れ物をここに置く
    (シム側は getattr で属性を読むだけなので実装差の影響を受けない)。"""
    def __init__(self, temperature=None, safety_settings=None,
                 system_instruction=None, response_mime_type=None, response_schema=None):
        self.temperature = temperature
        self.safety_settings = safety_settings
        self.system_instruction = system_instruction
        self.response_mime_type = response_mime_type
        self.response_schema = response_schema


class _CommonUsageMetadata:
    """response.usage_metadata 互換(トークン計測用)。本ツールは現時点で参照して
    いないが、他ツールのシムと契約を揃えておく。"""
    def __init__(self, usage):
        usage = usage if isinstance(usage, dict) else {}
        self.prompt_token_count = usage.get("promptTokenCount", 0)
        self.candidates_token_count = usage.get("candidatesTokenCount", 0)


class _CommonGeminiResponse:
    """client.models.generate_content(...) の戻り値互換。
    本ツールは response.parts で空応答を判定してから response.text を読むため、
    parts も提供する(空応答なら空リストになるので、呼び出し側の
    「空ならそのバッチは翻訳せず原文を返す」という既存の挙動がそのまま保たれる。
    本ツールにはリトライ機構が無いため、ここが原文のまま返る唯一の分岐になる)。
    レスポンスが想定外の形でも例外を投げず、text は空文字にする。"""
    def __init__(self, raw):
        try:
            self.text = raw["candidates"][0]["content"]["parts"][0]["text"]
        except (KeyError, IndexError, TypeError):
            self.text = ""
        try:
            self.parts = raw["candidates"][0]["content"]["parts"] or []
        except (KeyError, IndexError, TypeError):
            self.parts = []
        self.usage_metadata = _CommonUsageMetadata(
            raw.get("usageMetadata", {}) if isinstance(raw, dict) else {})


class _CommonGeminiModels:
    """client.models 互換。"""
    def generate_content(self, model=None, contents=None, config=None):
        if _generate_advanced is None:
            raise RuntimeError(_gemini_common_module_error_message())

        payload = {"contents": [{"parts": [{"text": contents}]}]}
        if config is not None:
            # safety_settings は旧SDKへ渡していた時点で既にREST形式
            # (category / threshold のdictリスト)なので、そのまま載せればよい。
            # 載せ忘れると BLOCK_NONE 指定が消え、資料の内容によっては応答が空になり
            # 「一部の段落だけ翻訳されていない」という切り分けにくい症状になる。
            safety = getattr(config, "safety_settings", None)
            if safety:
                payload["safetySettings"] = _schema_to_jsonable(safety)

            system_instruction = getattr(config, "system_instruction", None)
            if isinstance(system_instruction, str) and system_instruction:
                payload["systemInstruction"] = {"parts": [{"text": system_instruction}]}
            elif isinstance(system_instruction, dict):
                payload["systemInstruction"] = system_instruction

            gen_cfg = {}
            mime = getattr(config, "response_mime_type", None)
            if mime:
                gen_cfg["responseMimeType"] = mime
            schema = getattr(config, "response_schema", None)
            if schema is not None:
                gen_cfg["responseSchema"] = _schema_to_jsonable(schema)
            temp = getattr(config, "temperature", None)
            if temp is not None:
                gen_cfg["temperature"] = temp
            if gen_cfg:
                payload["generationConfig"] = gen_cfg

        # model は明示的に渡す(共通モジュール側の既定モデルへ勝手にフォールバック
        # されると、意図したモデルと実際に使われるモデルが食い違うため)。
        raw = _generate_advanced(payload, model=model)
        return _CommonGeminiResponse(raw)


class _CommonGeminiClient:
    """genai.Client(api_key=...) 相当。api_key は gemini_client.py 側が環境変数
    GEMINI_API_KEY から読むため、ここでは互換性のために受け取るだけで使用しない。"""
    def __init__(self, api_key=None):
        self.models = _CommonGeminiModels()


def gemini_credentials_available():
    """AI呼び出しが行える見込みがあるかどうかの事前チェック。
    直接呼び出しが遮断されていてもプロキシ経由なら成功しうるため、
    GEMINI_API_KEY / GEMINI_PROXY_URL のどちらか一方でも設定されていれば通す
    (プロキシ専用構成を誤って弾かないため)。"""
    return bool(os.environ.get("GEMINI_API_KEY") or os.environ.get("GEMINI_PROXY_URL"))


# --- 依存関係の確認 ---
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def check_dependencies(root_window):
    """起動時の依存関係チェック"""
    missing_libs = []
    if not HAS_DOCX:
        missing_libs.append("python-docx")
    
    if missing_libs:
        error_msg = "以下のライブラリがインストールされていません:\n"
        for lib in missing_libs:
            error_msg += f"- {lib}\n"
        error_msg += "\n以下のコマンドでインストールしてください:\n"
        error_msg += f"pip install {' '.join(missing_libs)}"
        
        messagebox.showerror("依存関係エラー", error_msg, parent=root_window)
        return False

    # 本ツールは全機能が翻訳(AI呼び出し)のため、共通モジュールが読めない場合は
    # 起動を続けても何もできない。原因が分かる形で案内して終了する。
    if not HAS_GEMINI:
        messagebox.showerror("依存関係エラー", _gemini_common_module_error_message(),
                             parent=root_window)
        return False

    return True

# --- グローバル変数（Gemini互換クライアント） ---
gemini_client = None

def init_gemini(root_window):
    """Gemini呼び出しの事前チェック。

    移行前はここで genai.configure() と genai.list_models() による自動モデル検出を
    行っていたが、list_models() はネットワークアクセスを伴うため、直接アクセスが
    遮断された環境では必ず例外になり「API初期化エラー」→ sys.exit(1) で
    ツールが起動すらできなかった。共通モジュール・プロキシのどちらにも list_models
    相当が無いため自動検出は廃止し、固定モデル名（環境変数 GEMINI_MODEL で上書き可）
    を使う方式へ変更した。この関数はネットワークへ一切アクセスしない。"""
    global gemini_client
    if _generate_advanced is None:
        messagebox.showerror("エラー", _gemini_common_module_error_message(),
                             parent=root_window)
        return False

    # 直接呼び出しが遮断されていてもプロキシ経由なら成功しうるため、
    # GEMINI_API_KEY 必須ではなく「どちらか一方でもあれば通す」に変更した。
    if not gemini_credentials_available():
        messagebox.showerror("エラー",
                             "Gemini認証情報が設定されていません。\n"
                             "以下のいずれかを設定してください:\n"
                             "- 環境変数 GEMINI_API_KEY （直接接続用）\n"
                             "- 環境変数 GEMINI_PROXY_URL （自宅PCプロキシ経由用）\n\n"
                             "※ setx で設定した場合は、コマンドプロンプトを\n"
                             "　 開き直してから起動してください。", parent=root_window)
        return False

    gemini_client = _CommonGeminiClient()
    print(f"使用モデル: {GEMINI_MODEL_NAME}")
    return True

def is_translatable(text):
    """翻訳が必要なテキストかどうかを判定"""
    if not text or str(text).strip() == "":
        return False
    
    text_str = str(text).strip()
    
    if text_str in ["", "#", "-", "N/A", "NULL", "•", "◦", "▪", "**", "*", ":", "：", 
                    "I.", "II.", "III.", "IV.", "V.", "VI.", "***"]:
        return False
    if text_str.replace(".", "").replace("-", "").isdigit():
        return False
    if len(text_str) <= 2:
        return False
    return True

def translate_batch_gemini(texts, target_language="Japanese"):
    """Gemini APIを使用した小バッチ翻訳"""
    if not gemini_client or not texts:
        return texts
    
    try:
        batch_input = "\n".join([f"[{i+1}] {t}" for i, t in enumerate(texts)])
        
        prompt = f"""
        Task: Translate the following text into {target_language}.
        
        Guidelines:
        1. Maintain the exact format [number] for each translated line.
        2. Output ONLY the numbered list. No extra explanations.
        3. Keep technical terms natural.
        
        Source Text:
        {batch_input}
        """

        safety_settings = [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
        ]

        # 旧: gemini_model.generate_content(
        #         prompt,
        #         generation_config=genai.types.GenerationConfig(temperature=0.1),
        #         safety_settings=safety_settings)
        # 新: 共通モジュール(gemini_client.py)経由の互換シムで同じ内容を送る。
        #     safety_settings は payload の safetySettings へそのまま載る。
        response = gemini_client.models.generate_content(
            model=GEMINI_MODEL_NAME,
            contents=prompt,
            config=_GeminiGenerateConfig(temperature=0.1,
                                         safety_settings=safety_settings),
        )
        
        time.sleep(1.5)

        if not response.parts:
            return texts

        response_text = response.text.strip()
        results = [None] * len(texts)
        lines = response_text.split('\n')
        
        for line in lines:
            match = re.match(r'^\[(\d+)\]\s*(.*)', line.strip())
            if match:
                idx = int(match.group(1)) - 1
                if 0 <= idx < len(texts):
                    results[idx] = match.group(2).strip()
        
        for i in range(len(results)):
            if results[i] is None:
                results[i] = texts[i]
                
        return results

    except Exception as e:
        print(f"バッチ翻訳エラー: {e}")
        time.sleep(2)
        return texts

def translate_super_fast_parallel(all_texts, target_language="Japanese", max_workers=3):
    """並列処理エンジン"""
    if not all_texts:
        return []
    
    batch_size = 10
    chunks = [all_texts[i:i + batch_size] for i in range(0, len(all_texts), batch_size)]
    results = [None] * len(chunks)
    
    def translate_chunk(chunk_idx, chunk_texts):
        translated = translate_batch_gemini(chunk_texts, target_language)
        return chunk_idx, translated

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = [executor.submit(translate_chunk, i, chunk) for i, chunk in enumerate(chunks)]
        
        for future in as_completed(futures):
            try:
                chunk_idx, translated_chunk = future.result()
                results[chunk_idx] = translated_chunk
            except Exception as e:
                print(f"チャンク処理エラー: {e}")
    
    final_results = []
    for chunk_result in results:
        if chunk_result:
            final_results.extend(chunk_result)
        else:
            final_results.extend([""] * batch_size)
            
    return final_results

class WordProgressWindow:
    """進捗表示用ウィンドウ（スレッドセーフ版）"""
    def __init__(self, parent):
        self.window = tk.Toplevel(parent)
        self.window.title("Gemini 翻訳進捗")
        self.window.geometry("450x180")
        self.window.resizable(False, False)
        
        try:
            self.window.transient(parent)
            self.window.grab_set()
        except:
            pass
        
        self.progress_label = tk.Label(self.window, text="Gemini AI 翻訳を準備中...", font=("Arial", 11, "bold"))
        self.progress_label.pack(pady=15)
        
        self.status_label = tk.Label(self.window, text="処理を開始します...", font=("Arial", 9))
        self.status_label.pack(pady=5)
        
        self.progress_frame = tk.Frame(self.window, width=350, height=20, bg="white", relief="sunken")
        self.progress_frame.pack(pady=10)
        
        self.progress_bar = tk.Frame(self.progress_frame, height=18, bg="#0078D4")
        self.progress_bar.place(x=1, y=1)
        
        self.time_label = tk.Label(self.window, text="", font=("Arial", 8), fg="blue")
        self.time_label.pack(pady=2)
        
        self.start_time = time.time()
        
    def update_progress(self, current, total, status=""):
        # 別スレッドから安全にGUIを更新するため after を使用
        try:
            self.window.after(0, self._update_gui, current, total, status)
        except Exception:
            pass
            
    def _update_gui(self, current, total, status):
        try:
            percentage = int((current / total) * 100) if total > 0 else 0
            bar_width = int((current / total) * 348) if total > 0 else 0
            self.progress_bar.config(width=bar_width)
            
            elapsed_time = time.time() - self.start_time
            
            self.progress_label.config(text=f"翻訳進捗: {current}/{total} ({percentage}%)")
            if status:
                self.status_label.config(text=status)
            self.time_label.config(text=f"経過時間: {elapsed_time:.1f}s")
        except:
            pass
        
    def close(self):
        try:
            self.window.after(0, self.window.destroy)
        except:
            pass

def translate_word_document_thread(file_path, target_language, progress_window):
    """バックグラウンドで実行されるメイン処理"""
    try:
        start_total_time = time.time()
        lang_code = target_language.split()[0].lower()
        output_path = os.path.splitext(file_path)[0] + f"_gemini_{lang_code}.docx"
        
        shutil.copy2(file_path, output_path)
        doc = Document(output_path)
        
        translatable_items = [] 
        for para in doc.paragraphs:
            for run in para.runs:
                if is_translatable(run.text):
                    translatable_items.append(run)
                    
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if is_translatable(run.text):
                                translatable_items.append(run)

        if not translatable_items:
            progress_window.close()
            messagebox.showinfo("完了", "翻訳対象のテキストが見つかりませんでした。")
            return

        progress_window.update_progress(0, len(translatable_items), f"Gemini APIで並列翻訳中... ({len(translatable_items)}項目)")
        
        texts_only = [run.text for run in translatable_items]
        
        # ※バックグラウンドスレッドで重い通信処理を実行
        translated_texts = translate_super_fast_parallel(texts_only, target_language, max_workers=3)
        
        progress_window.update_progress(len(translatable_items), len(translatable_items), "翻訳結果をWordに適用中...")
        for i, run in enumerate(translatable_items):
            if i < len(translated_texts) and translated_texts[i]:
                run.text = translated_texts[i]
                if "Japanese" in target_language or "日本" in target_language:
                    run.font.name = '游ゴシック'
        
        progress_window.update_progress(len(translatable_items), len(translatable_items), "保存中...")
        
        try:
            doc.save(output_path)
        except PermissionError:
            progress_window.close()
            messagebox.showerror("保存エラー", "ファイルが他のプログラム（Wordなど）で開かれています。\n閉じてから再度実行してください。")
            return

        progress_window.close()
        total_time = time.time() - start_total_time
        
        messagebox.showinfo("完了", 
                          f"書式保持翻訳完了！\n"
                          f"保存先: {output_path}\n"
                          f"翻訳項目数: {len(translatable_items)}\n"
                          f"処理時間: {total_time:.1f}秒")
        
    except Exception as e:
        progress_window.close()
        messagebox.showerror("エラー", f"翻訳処理中にエラーが発生しました:\n{str(e)}")

def select_file():
    path = filedialog.askopenfilename(
        title="翻訳するWordファイルを選択してください",
        filetypes=[("Word files", "*.docx"), ("Word files", "*.doc")]
    )
    
    if not path:
        return
    
    lang_win = tk.Toplevel(root)
    lang_win.title("Word翻訳設定")
    lang_win.geometry("400x250")
    lang_win.resizable(False, False)
    lang_win.transient(root)
    lang_win.grab_set()
    
    tk.Label(lang_win, text="翻訳先言語を選択してください", font=("Arial", 12, "bold")).pack(padx=20, pady=20)
    
    languages = {
        "日本語 (Japanese)": "Japanese",
        "英語 (English)": "English",
        "中国語簡体字 (Chinese)": "Chinese Simplified"
    }
    
    lang_var = tk.StringVar(lang_win)
    lang_var.set("日本語 (Japanese)")
    
    lang_menu = tk.OptionMenu(lang_win, lang_var, *languages.keys())
    lang_menu.config(font=("Arial", 10), width=25)
    lang_menu.pack(padx=20, pady=10)
    
    def start_translation():
        selected_language = languages[lang_var.get()]
        lang_win.destroy()
        
        # プログレスウィンドウを作成
        progress_window = WordProgressWindow(root)
        
        # 画面をフリーズさせないために、別スレッドで翻訳処理を開始！
        thread = threading.Thread(
            target=translate_word_document_thread, 
            args=(path, selected_language, progress_window)
        )
        thread.daemon = True
        thread.start()
    
    button_frame = tk.Frame(lang_win)
    button_frame.pack(pady=25)
    
    tk.Button(button_frame, text="翻訳開始", command=start_translation, 
             bg="#0078D4", fg="white", padx=20, pady=8, font=("Arial", 11, "bold")).pack(side=tk.LEFT, padx=10)
    tk.Button(button_frame, text="キャンセル", command=lang_win.destroy, 
             padx=20, pady=8, font=("Arial", 11)).pack(side=tk.LEFT, padx=10)

# --- GUI初期設定 ---
if __name__ == "__main__":
    root = tk.Tk()
    root.withdraw()
    
    if not check_dependencies(root):
        sys.exit(1)
        
    if not init_gemini(root):
        sys.exit(1)
        
    root.deiconify()
    root.title("Word Gemini 翻訳ツール")
    root.geometry("500x280")
    root.resizable(False, False)
    
    main_frame = tk.Frame(root)
    main_frame.pack(expand=True, fill='both', padx=20, pady=20)
    
    title_label = tk.Label(main_frame, text="Word Gemini 翻訳ツール", font=("Arial", 16, "bold"))
    title_label.pack(pady=8)
    
    subtitle_label = tk.Label(main_frame, text="書式完全保持版 (Gemini API)", font=("Arial", 12), fg="#0078D4")
    subtitle_label.pack(pady=2)
    
    desc_label = tk.Label(main_frame, 
                         text="Wordファイル(.docx)を選択して翻訳します\n"
                              "フォント、色、配置、テーブル書式を完全保持",
                         font=("Arial", 10))
    desc_label.pack(pady=8)
    
    select_button = tk.Button(main_frame, text="ファイル選択", command=select_file,
                             font=("Arial", 12), bg="#0078D4", fg="white", padx=20, pady=10)
    select_button.pack(pady=15)
    
    root.mainloop()