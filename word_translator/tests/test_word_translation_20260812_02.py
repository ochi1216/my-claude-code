"""word_translation_20260812_02.py の検証テスト。

20260812_02 は 20260812_01（Geminiプロキシ対応）へ、PowerPoint版が持っていた
次の5点を逆輸入した版である。テストもその5点の検証を主眼に置いている。

  1. 3回リトライ・3秒待機
  2. ロギング（translation_debug.log）
  3. バッチごとの進捗コールバック
  4. 3バッチ連続エラーでのフェイルファスト
  5. ファイルロックの事前検知（WinError 32）

あわせて、**Gemini通信経路（互換シム）が 20260812_01 から一切変わっていないこと**を
検証している（payload形状・safetySettings・モデル名・共通モジュール探索など）。

Windows / tkinter / 実際の Gemini API に依存せず検証するため、
偽の `gemini_client` を **対象モジュールのロード前に** sys.modules へ注入し、
`generate_advanced()` に渡る payload を捕捉して検証する。

python-docx は Linux コンテナにも入るため**本物を使う**。合成DOCX(見出し・本文
(同一段落に書式違いの複数run)・表、run ごとにフォントサイズ/太字/色を設定)を
実際に生成して翻訳文の書き戻しまでエンドツーエンドで検証し、run 単位の書式が
翻訳前後で変わっていないこと(＝書式保持を壊していないこと)を確認する。
さらに 20260812_01 と 20260812_02 に同じ翻訳文を与えて出力を比較し、
「機能追加でWord出力は何も変わっていない」ことを実物で示す。
tkinter だけスタブ化する。

実行方法:
    pip install python-docx
    python3 tests/test_word_translation_20260812_02.py
"""

import importlib.util
import inspect
import json
import os
import shutil
import socket
import sys
import tempfile
import time as _real_time
import types as pytypes

HERE = os.path.dirname(os.path.abspath(__file__))
TARGET = os.path.join(HERE, "..", "word_translation_20260812_02.py")
PREV_TARGET = os.path.join(HERE, "..", "word_translation_20260812_01.py")
OLD_TARGET = os.path.join(HERE, "..", "word_translation_20260306_01.py")

RESULTS = []


def check(name, ok, detail=""):
    RESULTS.append((name, bool(ok), detail))
    print(f"{'PASS' if ok else 'FAIL'}: {name}" + (f"  ({detail})" if detail else ""))


# ------------------------------------------------------------
# GUI 依存のスタブ (tkinter は apt が必要でコンテナに入らない)
# ------------------------------------------------------------
def _install_env_stubs():
    class _Anything:
        """属性アクセス・呼び出しを何でも受け付けるダミー。"""
        def __init__(self, *a, **k):
            pass

        def __call__(self, *a, **k):
            return _Anything()

        def __getattr__(self, _name):
            return _Anything()

    tk = pytypes.ModuleType("tkinter")
    for attr in ("Tk", "Toplevel", "Frame", "Label", "Button", "Entry", "Checkbutton",
                 "OptionMenu", "StringVar", "BooleanVar", "LEFT", "END"):
        setattr(tk, attr, _Anything())

    filedialog = pytypes.ModuleType("tkinter.filedialog")
    filedialog.askopenfilename = lambda *a, **k: ""

    # messagebox は「呼ばれたこと」を記録して、ガードの検証に使う
    messagebox = pytypes.ModuleType("tkinter.messagebox")
    messagebox.CALLS = []

    def _record(kind):
        def _fn(title="", message="", *a, **k):
            messagebox.CALLS.append((kind, title, message))
            return True
        return _fn

    messagebox.showerror = _record("error")
    messagebox.showwarning = _record("warning")
    messagebox.showinfo = _record("info")

    tk.filedialog = filedialog
    tk.messagebox = messagebox
    sys.modules["tkinter"] = tk
    sys.modules["tkinter.filedialog"] = filedialog
    sys.modules["tkinter.messagebox"] = messagebox
    return messagebox


MESSAGEBOX = _install_env_stubs()

from docx import Document  # noqa: E402  (スタブ注入後に読み込む)
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt, RGBColor  # noqa: E402


# ------------------------------------------------------------
# 偽 gemini_client
# ------------------------------------------------------------
CAPTURED = {"calls": []}
RESPONSE = {"value": None}


def _numbered_reply(payload):
    """プロンプト中の [n] の数だけ番号付きの訳文を返す。"""
    text = payload["contents"][0]["parts"][0]["text"]
    n = text.count("[")
    lines = "\n".join(f"[{i}] 訳文{i}" for i in range(1, n + 1))
    return {"candidates": [{"content": {"parts": [{"text": lines}]}}],
            "usageMetadata": {"promptTokenCount": 123, "candidatesTokenCount": 45}}


def _fake_generate_advanced(payload, model=None, **kwargs):
    CAPTURED["calls"].append({"payload": payload, "model": model})
    if RESPONSE["value"] is not None:
        value = RESPONSE["value"]
        if callable(value):
            return value(payload)
        return value
    return _numbered_reply(payload)


def _load_target(module_name, inject_fake=True, target=TARGET):
    """対象ファイルをロードする。inject_fake=False なら共通モジュール未配置を再現。"""
    sys.modules.pop("gemini_client", None)
    if inject_fake:
        fake = pytypes.ModuleType("gemini_client")
        fake.generate_advanced = _fake_generate_advanced
        sys.modules["gemini_client"] = fake

    spec = importlib.util.spec_from_file_location(module_name, target)
    mod = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = mod
    spec.loader.exec_module(mod)
    return mod


def _silence_sleep(mod):
    """成功後の待ち(1.5秒)・エラー後の待ち(2秒)でテストが遅くならないようにする。
    対象モジュールから見える time だけを差し替え、本物の time は壊さない。"""
    stub = pytypes.ModuleType("time")
    stub.sleep = lambda _s: None
    stub.time = __import__("time").time
    mod.time = stub


# ============================================================
# 検証
# ============================================================
mod = _load_target("target_main")
_silence_sleep(mod)

check("共通モジュールを読み込めた場合 HAS_GEMINI が True", mod.HAS_GEMINI is True)
check("旧SDK(google.generativeai)への依存が無い",
      not hasattr(mod, "genai") and "google.generativeai" not in sys.modules)
check("旧SDK由来のグローバル gemini_model が残っていない",
      not hasattr(mod, "gemini_model"))
check("python-docx への依存は残っている(HAS_DOCX)", mod.HAS_DOCX is True)


def _calls_named(path, name):
    """ソース中に「実際に呼び出している」箇所があるかをASTで調べる
    (コメント・docstringでの言及は除外する)。"""
    import ast
    tree = ast.parse(open(path, encoding="utf-8").read())
    for node in ast.walk(tree):
        if isinstance(node, ast.Call):
            fn = node.func
            if isinstance(fn, ast.Attribute) and fn.attr == name:
                return True
            if isinstance(fn, ast.Name) and fn.id == name:
                return True
    return False


check("自動モデル検出(list_models)の呼び出しが残っていない",
      not _calls_named(TARGET, "list_models"),
      "起動時のネットワークアクセスが無いこと")
check("genai.configure の呼び出しが残っていない", not _calls_named(TARGET, "configure"))
check("旧版には list_models の呼び出しがあった(移行の必要性の裏付け)",
      _calls_named(OLD_TARGET, "list_models"))

# --- init_gemini がネットワークアクセスを行わないこと -------------------
# 旧版は genai.list_models() を呼んでいたため、遮断下では必ず失敗して起動できなかった。
# ソケット生成そのものを禁止した状態で init_gemini() が成功することで、
# ネットワークアクセスが無くなったことを証明する。
os.environ["GEMINI_PROXY_URL"] = "https://example.invalid"
_real_socket = socket.socket


def _blocked_socket(*a, **k):
    raise AssertionError("init_gemini がネットワークアクセスを行った")


socket.socket = _blocked_socket
try:
    ok_init = mod.init_gemini(None)
except AssertionError as e:
    ok_init, _detail = False, str(e)
finally:
    socket.socket = _real_socket

check("init_gemini がネットワークアクセスなしで成功する", ok_init is True)
check("init_gemini がシムのクライアントを生成する",
      isinstance(mod.gemini_client, mod._CommonGeminiClient))
check("使用モデルが gemini-2.5-flash", mod.GEMINI_MODEL_NAME == "gemini-2.5-flash",
      f"model={mod.GEMINI_MODEL_NAME}")

# --- シグネチャに batch_idx / logger / progress_callback が入ったこと ----
_sig_batch = list(inspect.signature(mod.translate_batch_gemini).parameters)
check("translate_batch_gemini が batch_idx / logger を受け取る",
      _sig_batch == ["texts", "target_language", "batch_idx", "logger"],
      f"params={_sig_batch}")
_sig_par = list(inspect.signature(mod.translate_super_fast_parallel).parameters)
check("translate_super_fast_parallel が progress_callback / logger を受け取る",
      _sig_par == ["all_texts", "target_language", "max_workers",
                   "progress_callback", "logger"], f"params={_sig_par}")

# --- ロガー -------------------------------------------------------------
_log_dir = tempfile.mkdtemp()
_cwd_for_log = os.getcwd()
os.chdir(_log_dir)
try:
    _logger = mod.get_logger()
    _logger.info("テスト用の記録")
    for _h in _logger.handlers:
        _h.flush()
    _log_path = os.path.join(_log_dir, "translation_debug.log")
    check("get_logger が translation_debug.log を作る", os.path.isfile(_log_path),
          f"path={_log_path}")
    check("ログにメッセージが記録される",
          "テスト用の記録" in open(_log_path, encoding="utf-8").read())
    check("get_logger を2回呼んでもハンドラが増えない",
          mod.get_logger() is _logger and len(_logger.handlers) == 2,
          f"handlers={len(_logger.handlers)}")
finally:
    # ロガーは logging のグローバル辞書に残り、FileHandler はこの一時フォルダの
    # ファイルを掴んだままになる。後段のエンドツーエンド検証では別フォルダへ
    # 出力させたいので、ハンドラを片付けて get_logger() を作り直せる状態に戻す。
    for _h in list(_logger.handlers):
        _h.close()
        _logger.removeHandler(_h)
    os.chdir(_cwd_for_log)
    shutil.rmtree(_log_dir, ignore_errors=True)

# --- 通常の翻訳呼び出し -------------------------------------------------
CAPTURED["calls"].clear()
texts = ["Hello world", "This is a table cell"]
out, is_error = mod.translate_batch_gemini(texts, "Japanese", 0, None)

check("generate_advanced が1回呼ばれた(成功時はリトライしない)",
      len(CAPTURED["calls"]) == 1, f"calls={len(CAPTURED['calls'])}")
check("成功時はエラーフラグが False", is_error is False)
check("response.text の解析結果が返る", out == ["訳文1", "訳文2"], f"out={out}")

call = CAPTURED["calls"][0]
payload = call["payload"]

check("model が明示的に gemini-2.5-flash で渡る", call["model"] == "gemini-2.5-flash",
      f"model={call['model']!r}")
check("payload の contents が REST 形式",
      isinstance(payload.get("contents"), list)
      and isinstance(payload["contents"][0]["parts"][0]["text"], str))
check("payload に翻訳対象テキストが [番号] 付きで含まれる",
      "[1] Hello world" in payload["contents"][0]["parts"][0]["text"]
      and "[2] This is a table cell" in payload["contents"][0]["parts"][0]["text"])
check("翻訳先言語がプロンプトに反映される",
      "Japanese" in payload["contents"][0]["parts"][0]["text"])
check("generationConfig.temperature が 0.1 で camelCase で載る",
      payload.get("generationConfig", {}).get("temperature") == 0.1,
      f"generationConfig={payload.get('generationConfig')}")
check("payload が json.dumps 可能", bool(json.dumps(payload)))
check("payload に旧SDK固有の request_options が混ざっていない",
      "request_options" not in payload and "timeout" not in payload)

# --- safetySettings (載せ忘れると応答が空になりうる) --------------------
safety = payload.get("safetySettings")
check("safetySettings が payload に載る", isinstance(safety, list), f"safetySettings={safety}")
check("safetySettings が4カテゴリ分そのまま載る",
      isinstance(safety, list) and len(safety) == 4,
      f"len={len(safety) if isinstance(safety, list) else 'N/A'}")
check("safetySettings の全カテゴリが BLOCK_NONE",
      isinstance(safety, list)
      and all(s.get("threshold") == "BLOCK_NONE" for s in safety))
check("safetySettings のカテゴリ名が REST 形式(HARM_CATEGORY_*)",
      isinstance(safety, list)
      and {s.get("category") for s in safety} == {
          "HARM_CATEGORY_HARASSMENT", "HARM_CATEGORY_HATE_SPEECH",
          "HARM_CATEGORY_SEXUALLY_EXPLICIT", "HARM_CATEGORY_DANGEROUS_CONTENT"},
      f"categories={[s.get('category') for s in safety] if isinstance(safety, list) else safety}")

# --- レスポンス互換契約 (.parts / .text / usage_metadata) ---------------
resp = mod._CommonGeminiResponse(
    {"candidates": [{"content": {"parts": [{"text": "hello"}]}}],
     "usageMetadata": {"promptTokenCount": 7, "candidatesTokenCount": 3}})
check("response.text が読める", resp.text == "hello")
check("response.parts が読める(空応答判定に必須)",
      resp.parts == [{"text": "hello"}], f"parts={resp.parts}")
check("response.usage_metadata が読める",
      resp.usage_metadata.prompt_token_count == 7
      and resp.usage_metadata.candidates_token_count == 3)

for label, raw in [("空dict", {}), ("candidates空", {"candidates": []}),
                   ("None", None), ("想定外の型", {"candidates": "broken"}),
                   ("parts空", {"candidates": [{"content": {"parts": []}}]})]:
    try:
        r = mod._CommonGeminiResponse(raw)
        ok = r.text == "" and r.parts == [] and r.usage_metadata.prompt_token_count == 0
    except Exception as e:
        ok, label = False, f"{label}: {e}"
    check(f"壊れたレスポンス({label})で例外を投げず parts が [] になる", ok)

# --- 【20260812_02 の目玉1】空応答は3回リトライしてから原文へ ------------
# 20260812_01 までは空応答のとき1回で原文を返して終わっていた。
RESPONSE["value"] = {"candidates": [{"content": {"parts": []}}]}
CAPTURED["calls"].clear()
out2, is_error2 = mod.translate_batch_gemini(["原文A", "原文B"], "Japanese", 1, None)
check("空応答のとき3回リトライする", len(CAPTURED["calls"]) == 3,
      f"calls={len(CAPTURED['calls'])}")
check("空応答のとき最終的に原文を返す", out2 == ["原文A", "原文B"], f"out={out2}")
check("空応答のときエラーフラグが True", is_error2 is True)
RESPONSE["value"] = None


# --- 通信失敗も3回リトライしてから原文へ --------------------------------
def _raising(payload):
    raise RuntimeError("proxy down")


RESPONSE["value"] = _raising
CAPTURED["calls"].clear()
out3, is_error3 = mod.translate_batch_gemini(["原文A"], "Japanese", 2, None)
check("通信失敗時も3回リトライする", len(CAPTURED["calls"]) == 3,
      f"calls={len(CAPTURED['calls'])}")
check("通信失敗時は例外を投げず原文を返す", out3 == ["原文A"], f"out={out3}")
check("通信失敗時はエラーフラグが True", is_error3 is True)
RESPONSE["value"] = None

# --- 3回目で成功したらその結果を使う(リトライが機能している証明) ---------
_attempts = {"n": 0}


def _fail_twice(payload):
    _attempts["n"] += 1
    if _attempts["n"] < 3:
        raise RuntimeError("temporary failure")
    return _numbered_reply(payload)


RESPONSE["value"] = _fail_twice
CAPTURED["calls"].clear()
out3b, is_error3b = mod.translate_batch_gemini(["原文A"], "Japanese", 0, None)
check("2回失敗しても3回目に成功すれば訳文が返る", out3b == ["訳文1"], f"out={out3b}")
check("3回目で成功したらエラーフラグは False", is_error3b is False)
check("成功した時点でリトライを打ち切る", len(CAPTURED["calls"]) == 3,
      f"calls={len(CAPTURED['calls'])}")
RESPONSE["value"] = None

# --- リトライのログが記録されること -------------------------------------
class _RecordingLogger:
    def __init__(self):
        self.records = []

    def _rec(self, level):
        return lambda msg: self.records.append((level, str(msg)))

    def __getattr__(self, name):
        return self._rec(name)


RESPONSE["value"] = _raising
_rec_logger = _RecordingLogger()
mod.translate_batch_gemini(["原文A"], "Japanese", 4, _rec_logger)
check("リトライの試行回数がログに残る",
      sum(1 for lv, m in _rec_logger.records if "試行" in m) == 3,
      f"records={[m for _l, m in _rec_logger.records]}")
check("3回失敗したことがログに残る",
      any("3回失敗" in m for _lv, m in _rec_logger.records))
check("バッチ番号が1始まりでログに出る",
      any("バッチ 5" in m for _lv, m in _rec_logger.records),
      f"records={[m for _l, m in _rec_logger.records]}")
RESPONSE["value"] = None

# --- 番号が一部欠けた応答は該当項目だけ原文 -----------------------------
RESPONSE["value"] = {"candidates": [{"content": {"parts": [{"text": "[1] OK-1"}]}}]}
out4, _ = mod.translate_batch_gemini(["原文A", "原文B"], "Japanese", 3, None)
check("欠番の項目だけ原文にフォールバックする", out4 == ["OK-1", "原文B"], f"out={out4}")
RESPONSE["value"] = None

# --- 並列翻訳 -----------------------------------------------------------
res = mod.translate_super_fast_parallel(["text %d" % i for i in range(25)],
                                        "Japanese", max_workers=3)
check("並列翻訳が入力と同じ件数を返す", len(res) == 25, f"len={len(res)}")

CAPTURED["calls"].clear()
mod.translate_super_fast_parallel(["text %d" % i for i in range(25)], "Japanese",
                                  max_workers=3)
check("10件ずつのバッチに分割して送る(25件 -> 3バッチ)",
      len(CAPTURED["calls"]) == 3, f"calls={len(CAPTURED['calls'])}")

# --- 【20260812_02 の目玉2】進捗コールバック ----------------------------
_progress = []
mod.translate_super_fast_parallel(["text %d" % i for i in range(25)], "Japanese",
                                  max_workers=3, progress_callback=_progress.append)
check("進捗コールバックがバッチごとに呼ばれる", len(_progress) == 3,
      f"calls={_progress}")
check("進捗コールバックの最終値が全項目数になる",
      _progress and max(_progress) == 25,
      f"max={max(_progress) if _progress else None}")
check("進捗コールバックの値が単調増加する",
      _progress == sorted(_progress), f"progress={_progress}")
check("progress_callback を渡さなくても動く(既定 None)",
      len(mod.translate_super_fast_parallel(["text %d" % i for i in range(12)],
                                            "Japanese")) == 12)

# --- 【20260812_02 の目玉3】3バッチ連続エラーでフェイルファスト ----------
# 20260812_01 までは API 障害時も全バッチへリクエストを投げ続けていた。
#
# 注意: 中断は「ワーカーが次のバッチを取り出す前に abort_event が立つ」ことで効く。
# 通信もリトライ待ちも即座に返るテスト環境では、ワーカーが全バッチを走り終えてから
# メインスレッドが3件目のエラーを数えることになり、中断が間に合わない。
# 実機の通信時間を模して1呼び出しあたり実時間で少し待たせる。
def _raising_slow(payload):
    _real_time.sleep(0.02)
    raise RuntimeError("proxy down")


RESPONSE["value"] = _raising_slow
CAPTURED["calls"].clear()
try:
    mod.translate_super_fast_parallel(["a" * 5] * 100, "Japanese", max_workers=1)
    check("3バッチ連続エラーでフェイルファストする", False, "RuntimeErrorが出なかった")
except RuntimeError as e:
    check("3バッチ連続エラーでフェイルファストする", True)
    check("フェイルファストのメッセージが原因を説明している",
          "3回連続で失敗" in str(e), str(e))
    # 10バッチ全部に投げていたら 10バッチ x 3リトライ = 30回になる。
    # 3バッチ目で撤退するので、それより十分少ないことを確認する。
    check("フェイルファストで残りのバッチへは投げない",
          len(CAPTURED["calls"]) < 30, f"calls={len(CAPTURED['calls'])} (投げ切りなら30)")
RESPONSE["value"] = None

# --- 散発的な失敗はリトライで拾われ、フェイルファストに至らないこと ------
_seq = {"n": 0}


def _flaky(payload):
    """1回おきに失敗させる。各バッチは1回失敗しても再試行で成功するため、
    バッチ単位のエラーは発生せず、連続エラーカウントは増えない。"""
    _seq["n"] += 1
    if _seq["n"] % 2 == 1:
        raise RuntimeError("intermittent")
    return _numbered_reply(payload)


RESPONSE["value"] = _flaky
try:
    _flaky_out = mod.translate_super_fast_parallel(["text %d" % i for i in range(50)],
                                                   "Japanese", max_workers=1)
    check("1回失敗しても再試行で成功すればフェイルファストしない",
          len(_flaky_out) == 50, f"len={len(_flaky_out)}")
    check("再試行で拾えたバッチは訳文が返る",
          all(t.startswith("訳文") for t in _flaky_out),
          f"sample={_flaky_out[:3]}")
except RuntimeError as e:
    check("1回失敗しても再試行で成功すればフェイルファストしない", False, str(e))
RESPONSE["value"] = None

check("空リストを渡すと空リストが返る",
      mod.translate_super_fast_parallel([], "Japanese") == [])

# --- is_translatable (移行と無関係だが回帰確認) -------------------------
check("is_translatable: 通常の文は翻訳対象", mod.is_translatable("Hello world") is True)
check("is_translatable: 空文字は対象外", mod.is_translatable("") is False)
check("is_translatable: 2文字以下は対象外", mod.is_translatable("ab") is False)
check("is_translatable: 数字だけは対象外", mod.is_translatable("12.34") is False)
check("is_translatable: 箇条書き記号は対象外", mod.is_translatable("•") is False)

# --- 認証情報の判定 (プロキシURLのみ = 通ることが最重要) ----------------
_env_backup = {k: os.environ.get(k) for k in ("GEMINI_API_KEY", "GEMINI_PROXY_URL")}
for api, proxy, expected, label in [
        ("k", None, True, "APIキーのみ"),
        (None, "https://x", True, "プロキシURLのみ(プロキシ専用構成)"),
        ("k", "https://x", True, "両方"),
        (None, None, False, "どちらも未設定")]:
    for name, val in (("GEMINI_API_KEY", api), ("GEMINI_PROXY_URL", proxy)):
        os.environ.pop(name, None)
        if val:
            os.environ[name] = val
    check(f"gemini_credentials_available: {label} -> {expected}",
          mod.gemini_credentials_available() is expected)

    MESSAGEBOX.CALLS.clear()
    check(f"init_gemini: {label} -> {expected}", mod.init_gemini(None) is expected)
    if not expected:
        msg = MESSAGEBOX.CALLS[-1][2] if MESSAGEBOX.CALLS else ""
        check("未設定時のエラーが両方の環境変数を案内する",
              "GEMINI_API_KEY" in msg and "GEMINI_PROXY_URL" in msg, msg)

for k, v in _env_backup.items():
    os.environ.pop(k, None)
    if v:
        os.environ[k] = v

# --- 共通モジュール未配置時 ---------------------------------------------
mod2 = _load_target("target_no_common", inject_fake=False)
check("共通モジュール未配置なら HAS_GEMINI が False", mod2.HAS_GEMINI is False)

try:
    mod2._CommonGeminiClient().models.generate_content(model="m", contents="c")
    check("共通モジュール未配置でAI呼び出しすると RuntimeError", False, "例外が出なかった")
except RuntimeError as e:
    text = str(e)
    check("共通モジュール未配置でAI呼び出しすると RuntimeError", True)
    check("エラーに探索したパスが含まれる", "common" in text, text.splitlines()[1])
    check("エラーに GEMINI_COMMON_DIR の案内が含まれる", "GEMINI_COMMON_DIR" in text)
except Exception as e:
    check("共通モジュール未配置でAI呼び出しすると RuntimeError", False, repr(e))

MESSAGEBOX.CALLS.clear()
check("共通モジュール未配置なら check_dependencies が False",
      mod2.check_dependencies(None) is False)
check("check_dependencies のエラーが原因を説明している",
      any("gemini_client.py" in c[2] for c in MESSAGEBOX.CALLS),
      f"calls={MESSAGEBOX.CALLS}")
MESSAGEBOX.CALLS.clear()
check("共通モジュール未配置なら init_gemini も False", mod2.init_gemini(None) is False)
check("依存関係が揃っていれば check_dependencies が True",
      mod.check_dependencies(None) is True)

# --- 共通モジュールの探索 (会社PCは「2つ上」が common) ------------------
_tmp = tempfile.mkdtemp()
try:
    _common = os.path.join(_tmp, "PythonScripts", "common")
    _tool = os.path.join(_tmp, "PythonScripts", "word", "word_translator")
    os.makedirs(_common)
    os.makedirs(_tool)
    with open(os.path.join(_common, "gemini_client.py"), "w", encoding="utf-8") as f:
        f.write("MARKER = 'two-up'\n\n\ndef generate_advanced(payload, model=None):\n"
                "    return {'candidates': [{'content': {'parts': [{'text': '[1] ok'}]}}]}\n")
    _copied = os.path.join(_tool, os.path.basename(TARGET))
    shutil.copy(TARGET, _copied)

    _saved_path = list(sys.path)
    mod3 = _load_target("target_two_up", inject_fake=False, target=_copied)
    check("「2つ上が common」の配置(会社PCの実配置)で共通モジュールを解決できる",
          mod3.HAS_GEMINI is True and mod3._COMMON_DIR == os.path.abspath(_common),
          f"COMMON_DIR={getattr(mod3, '_COMMON_DIR', None)}")
    check("探索候補に1つ上・2つ上・3つ上が含まれる",
          len(mod3._COMMON_DIR_CANDIDATES) == 3, f"{mod3._COMMON_DIR_CANDIDATES}")

    # 「1つ上が common」の配置(他ツールと同じ置き方)でも解決できること
    _flat = os.path.join(_tmp, "PythonScripts", "word")
    _copied2 = os.path.join(_flat, os.path.basename(TARGET))
    shutil.copy(TARGET, _copied2)
    mod3b = _load_target("target_one_up", inject_fake=False, target=_copied2)
    check("「1つ上が common」の配置でも解決できる",
          mod3b.HAS_GEMINI is True and mod3b._COMMON_DIR == os.path.abspath(_common),
          f"COMMON_DIR={getattr(mod3b, '_COMMON_DIR', None)}")

    # GEMINI_COMMON_DIR による明示指定
    os.environ["GEMINI_COMMON_DIR"] = _common
    mod4 = _load_target("target_env_dir", inject_fake=False)
    check("GEMINI_COMMON_DIR で共通モジュールの場所を明示指定できる",
          mod4.HAS_GEMINI is True
          and mod4._COMMON_DIR_CANDIDATES == [_common],
          f"candidates={mod4._COMMON_DIR_CANDIDATES}")
    os.environ.pop("GEMINI_COMMON_DIR", None)
    sys.path[:] = _saved_path
finally:
    shutil.rmtree(_tmp, ignore_errors=True)
    sys.modules.pop("gemini_client", None)

# --- 未使用オプション(他ツールとの契約統一) -----------------------------
CAPTURED["calls"].clear()
sys.modules["gemini_client"] = pytypes.ModuleType("gemini_client")
cfg = mod._GeminiGenerateConfig(temperature=0.5, response_mime_type="application/json",
                                response_schema={"type": "object"},
                                system_instruction="be brief")
mod._CommonGeminiClient().models.generate_content(model="m", contents="c", config=cfg)
_p = CAPTURED["calls"][-1]["payload"]
check("responseMimeType / responseSchema が camelCase で載る",
      _p.get("generationConfig", {}).get("responseMimeType") == "application/json"
      and _p.get("generationConfig", {}).get("responseSchema") == {"type": "object"},
      f"generationConfig={_p.get('generationConfig')}")
check("systemInstruction が REST 形式で載る",
      _p.get("systemInstruction") == {"parts": [{"text": "be brief"}]},
      f"systemInstruction={_p.get('systemInstruction')}")

CAPTURED["calls"].clear()
mod._CommonGeminiClient().models.generate_content(model="m", contents="c", config=None)
check("config が None でも payload を組み立てられる",
      "generationConfig" not in CAPTURED["calls"][-1]["payload"]
      and "safetySettings" not in CAPTURED["calls"][-1]["payload"])


class _FakePydanticSchema:
    def model_dump(self, **kwargs):
        return {"type": "OBJECT", "propertyOrdering": ["a"]}


check("pydantic モデルの response_schema も dict 化できる",
      mod._schema_to_jsonable(_FakePydanticSchema()) == {"type": "OBJECT",
                                                         "propertyOrdering": ["a"]})


# ============================================================
# Word処理のエンドツーエンド検証 (python-docx は本物を使う)
#   移行で Word の書式保持処理に一切影響が出ていないことを、合成DOCXで確認する。
#   run 単位の書式(サイズ/太字/色/段落配置/段落スタイル)が翻訳前後で一致すれば、
#   「書式完全保持」を壊していない証拠になる。
# ============================================================
class _FakeProgressWindow:
    """WordProgressWindow 互換の記録用ダミー(tkinter を使わない)。"""
    def __init__(self):
        self.updates = []
        self.closed = 0

    def update_progress(self, current, total, status=""):
        self.updates.append((current, total, status))

    def close(self):
        self.closed += 1


def _build_sample_docx(path):
    """見出し ＋ 本文(同一段落に書式違いの複数run) ＋ 表 を含むDOCXを作る。"""
    doc = Document()

    # --- 見出し ---
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("Employee Engagement Survey Results")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x30, 0x30)

    # --- 本文: 1つの段落に書式の違う run を2つ置く(run単位保持の検証に使う) ---
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run("This paragraph describes the overall findings.")
    r1.font.size = Pt(11)
    r1.font.bold = False
    r1.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    r2 = p1.add_run(" Response rate reached a record high.")
    r2.font.size = Pt(9)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p2.add_run("Prepared by the strategy office.")
    r3.font.size = Pt(8)
    r3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # --- 翻訳対象外になるはずの run (2文字以下 / 記号 / 数字) ---
    p3 = doc.add_paragraph()
    for txt in ("ab", "•", "12.34"):
        rr = p3.add_run(txt)
        rr.font.size = Pt(10)

    # --- 表 ---
    table = doc.add_table(rows=3, cols=2)
    cells = [("Category", "Score"), ("Wellbeing", "Sixty six percent"),
             ("Leadership", "Seventy two percent")]
    for i, (left, right) in enumerate(cells):
        for j, txt in enumerate((left, right)):
            para = table.cell(i, j).paragraphs[0]
            run = para.add_run(txt)
            run.font.size = Pt(10 + i)
            run.font.bold = (i == 0)
            run.font.color.rgb = RGBColor(0x11, 0x22, 0x33)

    doc.save(path)


def _run_signature(path):
    """run 単位のテキストと書式を、比較できる形で書き出す。
    翻訳対象と同じ経路(doc.paragraphs / doc.tables)を辿る。"""
    doc = Document(path)
    sig = []

    def _walk(para, key):
        for ri, run in enumerate(para.runs):
            try:
                color = str(run.font.color.rgb)
            except Exception:
                color = None
            sig.append({
                "key": key + (ri,),
                "text": run.text,
                "size": run.font.size.pt if run.font.size is not None else None,
                "bold": run.font.bold,
                "name": run.font.name,
                "color": color,
                "align": str(para.alignment),
                "style": para.style.name if para.style is not None else None,
            })

    for pi, para in enumerate(doc.paragraphs):
        _walk(para, ("body", pi))
    for ti, table in enumerate(doc.tables):
        for ri_, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, para in enumerate(cell.paragraphs):
                    _walk(para, ("table", ti, ri_, ci, pi))
    return sig


def _fmt_only(sig):
    """書式だけを抜き出す(テキストとフォント名は翻訳で変わるので除外)。"""
    return [{k: v for k, v in item.items() if k not in ("text", "name")} for item in sig]


_docx_dir = tempfile.mkdtemp()
_cwd = os.getcwd()
try:
    # translation_debug.log がリポジトリに落ちないよう、一時フォルダで実行する
    os.chdir(_docx_dir)

    src_docx = os.path.join(_docx_dir, "sample.docx")
    _build_sample_docx(src_docx)
    before_sig = _run_signature(src_docx)

    CAPTURED["calls"].clear()
    sys.modules["gemini_client"] = pytypes.ModuleType("gemini_client")
    pw = _FakeProgressWindow()
    MESSAGEBOX.CALLS.clear()
    mod.translate_word_document_thread(src_docx, "Japanese", pw)

    out_docx = os.path.join(_docx_dir, "sample_gemini_japanese.docx")
    check("出力ファイル名が 元ファイル名_gemini_japanese.docx になる",
          os.path.isfile(out_docx), f"exists={os.path.isfile(out_docx)}")
    check("元のファイルは書き換えられていない",
          _run_signature(src_docx) == before_sig)
    check("エラーダイアログが出ていない",
          not any(c[0] == "error" for c in MESSAGEBOX.CALLS), f"{MESSAGEBOX.CALLS}")
    check("完了ダイアログが出る", any(c[0] == "info" for c in MESSAGEBOX.CALLS))
    check("進捗ウィンドウが閉じられる", pw.closed >= 1, f"closed={pw.closed}")
    check("進捗ウィンドウが更新される", len(pw.updates) > 0, f"updates={len(pw.updates)}")
    check("翻訳がシム経由(共通モジュール)で行われた", len(CAPTURED["calls"]) >= 1,
          f"calls={len(CAPTURED['calls'])}")

    # 20260812_02: 翻訳中にも進捗が進むこと(0 と 完了 の2段階で終わらないこと)
    _mid = [u for u in pw.updates if "並列翻訳中" in u[2]]
    check("翻訳中の進捗更新が「開始」と「バッチ完了」の2回以上ある",
          len(_mid) >= 2, f"updates={[u[2] for u in pw.updates]}")
    check("進捗表示が 現在値/全体 の形式になる",
          any(f"({u[0]}/{u[1]}項目)" in u[2] for u in _mid),
          f"updates={[u[2] for u in _mid]}")

    # 20260812_02: ログファイルが作られ、開始と完了が記録されること
    _e2e_log = os.path.join(_docx_dir, "translation_debug.log")
    for _h in mod.get_logger().handlers:
        _h.flush()
    check("翻訳を実行すると translation_debug.log が作られる",
          os.path.isfile(_e2e_log), f"path={_e2e_log}")
    _log_text = open(_e2e_log, encoding="utf-8").read() if os.path.isfile(_e2e_log) else ""
    check("ログに翻訳開始が記録される", "Word翻訳開始" in _log_text)
    check("ログに処理完了が記録される", "処理完了" in _log_text)
    check("ログにバッチの成功が記録される", "成功" in _log_text)
    check("エンドツーエンドでも model が明示的に渡る",
          all(c["model"] == "gemini-2.5-flash" for c in CAPTURED["calls"]))
    check("エンドツーエンドでも safetySettings が載る",
          all(len(c["payload"].get("safetySettings", [])) == 4 for c in CAPTURED["calls"]))

    after_sig = _run_signature(out_docx)

    check("run の構成(数と位置)が翻訳前後で変わらない",
          [s["key"] for s in before_sig] == [s["key"] for s in after_sig],
          f"before={len(before_sig)} / after={len(after_sig)}")
    check("run 単位の書式(サイズ/太字/色/段落配置/段落スタイル)が翻訳前後で完全一致する",
          _fmt_only(before_sig) == _fmt_only(after_sig),
          "差分あり" if _fmt_only(before_sig) != _fmt_only(after_sig) else "")

    _translated = [s for s in after_sig if s["text"].startswith("訳文")]
    check("翻訳文が run に書き戻されている", len(_translated) > 0,
          f"translated={len(_translated)}")
    check("日本語指定時に run.font.name が 游ゴシック になる",
          all(s["name"] == "游ゴシック" for s in _translated),
          f"names={sorted({s['name'] for s in _translated})}")

    # 表のセルも翻訳対象として拾えていること (doc.tables 経路)
    check("表のセルが翻訳対象として拾われている",
          any(s["key"][0] == "table" and s["text"].startswith("訳文") for s in after_sig),
          f"table_runs={[s['key'] for s in after_sig if s['key'][0] == 'table']}")
    check("本文の段落が翻訳対象として拾われている",
          any(s["key"][0] == "body" and s["text"].startswith("訳文") for s in after_sig))

    # 翻訳対象外のものは原文のまま
    _texts = [s["text"] for s in after_sig]
    check("2文字以下・記号・数字だけの run は原文のまま",
          all(t in _texts for t in ("ab", "•", "12.34")), f"texts={_texts}")

    # --- 20260812_01 と 20260812_02 でDOCX出力が一致することの直接確認 -----
    # 今回追加したのはリトライ・ログ・進捗・フェイルファスト・ロック検知だけで、
    # Word の書き戻し処理そのものは触っていない。同じ翻訳文を与えれば
    # 出力DOCX(テキスト・書式)は完全に一致するはずである。
    prev_mod = _load_target("target_prev", target=PREV_TARGET)
    _silence_sleep(prev_mod)
    prev_mod.gemini_client = prev_mod._CommonGeminiClient()

    prev_src = os.path.join(_docx_dir, "sample_prev.docx")
    shutil.copy(src_docx, prev_src)
    prev_mod.translate_word_document_thread(prev_src, "Japanese", _FakeProgressWindow())
    prev_out = os.path.join(_docx_dir, "sample_prev_gemini_japanese.docx")

    check("20260812_01 でも同じ名前で出力される", os.path.isfile(prev_out))
    prev_sig = _run_signature(prev_out)
    check("20260812_01 と 20260812_02 で出力DOCXの run 構成が一致する",
          [s["key"] for s in prev_sig] == [s["key"] for s in after_sig],
          f"prev={len(prev_sig)} / new={len(after_sig)}")
    check("20260812_01 と 20260812_02 で出力DOCXのテキストが完全一致する",
          [s["text"] for s in prev_sig] == [s["text"] for s in after_sig])
    check("20260812_01 と 20260812_02 で出力DOCXの書式が完全一致する",
          _fmt_only(prev_sig) == _fmt_only(after_sig))
    check("20260812_01 と 20260812_02 で出力DOCXのフォント名が完全一致する",
          [s["name"] for s in prev_sig] == [s["name"] for s in after_sig])

    # 偽 gemini_client を注入し直す(prev_mod のロードで差し替わっているため)
    sys.modules["gemini_client"] = pytypes.ModuleType("gemini_client")

    # --- 移行前(_20260306_01)と比べても出力が変わっていないこと ------------
    # ※DOCX(zip)はタイムスタンプが毎回変わるためバイト比較はできない。
    #   run 単位のテキストと書式で比較する。
    if os.path.isfile(OLD_TARGET):
        class _OldResponse:
            def __init__(self, payload_text):
                n = payload_text.count("[")
                self.text = "\n".join(f"[{i}] 訳文{i}" for i in range(1, n + 1))
                self.parts = [{"text": self.text}]

        class _OldModel:
            def generate_content(self, prompt, **kwargs):
                return _OldResponse(prompt)

        fake_genai = pytypes.ModuleType("google.generativeai")
        fake_genai.configure = lambda **k: None
        fake_genai.list_models = lambda: []
        fake_genai.GenerativeModel = lambda *a, **k: _OldModel()
        fake_genai.types = pytypes.SimpleNamespace(
            GenerationConfig=lambda **k: pytypes.SimpleNamespace(**k))
        google_pkg = pytypes.ModuleType("google")
        google_pkg.generativeai = fake_genai
        sys.modules["google"] = google_pkg
        sys.modules["google.generativeai"] = fake_genai

        spec = importlib.util.spec_from_file_location("target_old", OLD_TARGET)
        old_mod = importlib.util.module_from_spec(spec)
        sys.modules["target_old"] = old_mod
        spec.loader.exec_module(old_mod)
        _silence_sleep(old_mod)
        old_mod.gemini_model = _OldModel()

        old_src = os.path.join(_docx_dir, "sample_old.docx")
        shutil.copy(src_docx, old_src)
        old_mod.translate_word_document_thread(old_src, "Japanese", _FakeProgressWindow())
        old_out = os.path.join(_docx_dir, "sample_old_gemini_japanese.docx")

        check("移行前(20260306_01)でも同じ名前で出力される", os.path.isfile(old_out))
        old_sig = _run_signature(old_out)
        check("移行前と 20260812_02 で出力DOCXの run 構成が一致する",
              [s["key"] for s in old_sig] == [s["key"] for s in after_sig],
              f"old={len(old_sig)} / new={len(after_sig)}")
        check("移行前と 20260812_02 で出力DOCXのテキストが完全一致する",
              [s["text"] for s in old_sig] == [s["text"] for s in after_sig])
        check("移行前と 20260812_02 で出力DOCXの書式が完全一致する",
              _fmt_only(old_sig) == _fmt_only(after_sig))
        check("移行前と 20260812_02 で出力DOCXのフォント名が完全一致する",
              [s["name"] for s in old_sig] == [s["name"] for s in after_sig])
    else:
        check("移行前(_20260306_01)が見つかり比較できる", False, f"not found: {OLD_TARGET}")

    # --- 翻訳対象が無いDOCXでも落ちないこと -------------------------------
    empty_docx = os.path.join(_docx_dir, "empty.docx")
    Document().save(empty_docx)
    MESSAGEBOX.CALLS.clear()
    pw2 = _FakeProgressWindow()
    mod.translate_word_document_thread(empty_docx, "Japanese", pw2)
    check("翻訳対象が無いDOCXでもエラーにならず案内が出る",
          any(c[0] == "info" for c in MESSAGEBOX.CALLS)
          and not any(c[0] == "error" for c in MESSAGEBOX.CALLS),
          f"{MESSAGEBOX.CALLS}")

    # --- ヘッダー/フッターが翻訳対象外であること(既存仕様。移行で広げていない) ---
    hf_docx = os.path.join(_docx_dir, "header_footer.docx")
    _hf = Document()
    _hf.add_paragraph("Body text that should be translated.")
    _hf.sections[0].header.paragraphs[0].text = "Header text stays in English."
    _hf.sections[0].footer.paragraphs[0].text = "Footer text stays in English."
    _hf.save(hf_docx)
    MESSAGEBOX.CALLS.clear()
    mod.translate_word_document_thread(hf_docx, "Japanese", _FakeProgressWindow())
    _hf_out = Document(os.path.join(_docx_dir, "header_footer_gemini_japanese.docx"))
    check("ヘッダー/フッターは翻訳対象外のまま(既存仕様。移行で範囲を広げていない)",
          _hf_out.sections[0].header.paragraphs[0].text == "Header text stays in English."
          and _hf_out.sections[0].footer.paragraphs[0].text == "Footer text stays in English.")
    check("同じ文書の本文は翻訳されている",
          _hf_out.paragraphs[0].runs[0].text.startswith("訳文"),
          f"body={_hf_out.paragraphs[0].text!r}")

    # --- 【20260812_02 の目玉4】ファイルロックの事前検知 -------------------
    # 20260812_01 までは数分かけた翻訳が終わった後、保存時に初めて失敗していた。
    # 実機の WinError 32 を Linux で再現できないため、対象モジュールから見える
    # open() だけを差し替えて「追記で開けないファイル」を作る。
    import builtins

    def _make_locked_open(locked_path):
        def _fake_open(path, mode="r", *a, **k):
            if os.path.abspath(path) == os.path.abspath(locked_path) and "a" in mode:
                raise PermissionError(32, "The process cannot access the file")
            return builtins.open(path, mode, *a, **k)
        return _fake_open

    lock_src = os.path.join(_docx_dir, "locked.docx")
    _build_sample_docx(lock_src)

    # (1) 読み込み元がロックされている場合
    CAPTURED["calls"].clear()
    MESSAGEBOX.CALLS.clear()
    pw3 = _FakeProgressWindow()
    mod.open = _make_locked_open(lock_src)
    try:
        mod.translate_word_document_thread(lock_src, "Japanese", pw3)
    finally:
        del mod.open
    check("読み込み元がロックされていると翻訳前に中止する",
          any(c[0] == "error" for c in MESSAGEBOX.CALLS), f"{MESSAGEBOX.CALLS}")
    check("ロック検知のメッセージがファイルを閉じるよう案内する",
          any("開かれています" in c[2] for c in MESSAGEBOX.CALLS),
          f"{MESSAGEBOX.CALLS}")
    check("ロック検知時はAPIを一度も呼ばない(数分待たされない)",
          len(CAPTURED["calls"]) == 0, f"calls={len(CAPTURED['calls'])}")
    check("ロック検知時も進捗ウィンドウは閉じる", pw3.closed >= 1, f"closed={pw3.closed}")
    check("ロック検知時は出力ファイルを作らない",
          not os.path.isfile(os.path.join(_docx_dir, "locked_gemini_japanese.docx")))

    # (2) 保存先(前回の出力)がロックされている場合
    lock_src2 = os.path.join(_docx_dir, "locked2.docx")
    _build_sample_docx(lock_src2)
    lock_out2 = os.path.join(_docx_dir, "locked2_gemini_japanese.docx")
    shutil.copy(lock_src2, lock_out2)          # 前回の出力が残っている状態を作る
    CAPTURED["calls"].clear()
    MESSAGEBOX.CALLS.clear()
    pw4 = _FakeProgressWindow()
    mod.open = _make_locked_open(lock_out2)
    try:
        mod.translate_word_document_thread(lock_src2, "Japanese", pw4)
    finally:
        del mod.open
    check("保存先がロックされていると翻訳前に中止する",
          any(c[0] == "error" for c in MESSAGEBOX.CALLS), f"{MESSAGEBOX.CALLS}")
    check("保存先ロックのメッセージが以前の出力ファイルを指す",
          any("以前に作成した翻訳ファイル" in c[2] for c in MESSAGEBOX.CALLS),
          f"{MESSAGEBOX.CALLS}")
    check("保存先ロック検知時もAPIを一度も呼ばない",
          len(CAPTURED["calls"]) == 0, f"calls={len(CAPTURED['calls'])}")

    # ロックが無ければ従来どおり翻訳できること(事前チェックが誤検知しない)
    MESSAGEBOX.CALLS.clear()
    mod.translate_word_document_thread(lock_src2, "Japanese", _FakeProgressWindow())
    check("ロックが無ければ事前チェックを素通りして翻訳できる",
          any(c[0] == "info" for c in MESSAGEBOX.CALLS)
          and not any(c[0] == "error" for c in MESSAGEBOX.CALLS), f"{MESSAGEBOX.CALLS}")

    # --- フェイルファストがGUI側まで伝わること -----------------------------
    ff_src = os.path.join(_docx_dir, "failfast.docx")
    _ff = Document()
    for i in range(60):
        _ff.add_paragraph(f"Sentence number {i} to be translated.")
    _ff.save(ff_src)
    RESPONSE["value"] = _raising_slow
    MESSAGEBOX.CALLS.clear()
    pw5 = _FakeProgressWindow()
    mod.translate_word_document_thread(ff_src, "Japanese", pw5)
    check("フェイルファストが「通信エラー強制終了」ダイアログになる",
          any(c[1] == "通信エラー強制終了" for c in MESSAGEBOX.CALLS),
          f"{MESSAGEBOX.CALLS}")
    check("フェイルファスト時も進捗ウィンドウは閉じる", pw5.closed >= 1,
          f"closed={pw5.closed}")
    # 中断されると doc.save() まで到達しないため、出力先には「翻訳前の複製」が
    # 残る（PowerPoint版と同じ挙動）。中途半端に翻訳された文書にはならない。
    _ff_out = os.path.join(_docx_dir, "failfast_gemini_japanese.docx")
    check("フェイルファスト時は中途半端に翻訳された文書を保存しない",
          not os.path.isfile(_ff_out)
          or [p.text for p in Document(_ff_out).paragraphs]
          == [p.text for p in Document(ff_src).paragraphs],
          "出力に翻訳結果が混ざっている")
    RESPONSE["value"] = None

    # --- 予期せぬエラーで traceback がログに残ること ------------------------
    for _h in mod.get_logger().handlers:
        _h.flush()
    _before_len = len(open(_e2e_log, encoding="utf-8").read())
    MESSAGEBOX.CALLS.clear()
    mod.translate_word_document_thread(os.path.join(_docx_dir, "not_exist.docx"),
                                       "Japanese", _FakeProgressWindow())
    for _h in mod.get_logger().handlers:
        _h.flush()
    _tail = open(_e2e_log, encoding="utf-8").read()[_before_len:]
    check("予期せぬエラーでダイアログが出る",
          any(c[0] == "error" for c in MESSAGEBOX.CALLS), f"{MESSAGEBOX.CALLS}")
    check("予期せぬエラーの traceback がログに残る",
          "Traceback" in _tail, f"tail={_tail[:200]}")
finally:
    os.chdir(_cwd)
    shutil.rmtree(_docx_dir, ignore_errors=True)
    sys.modules.pop("google", None)
    sys.modules.pop("google.generativeai", None)


# ============================================================
print("\n" + "=" * 60)
failed = [r for r in RESULTS if not r[1]]
print(f"合計 {len(RESULTS)} 項目 / 合格 {len(RESULTS) - len(failed)} / 失敗 {len(failed)}")
if failed:
    for name, _ok, detail in failed:
        print(f"  FAILED: {name}  {detail}")
sys.exit(1 if failed else 0)
